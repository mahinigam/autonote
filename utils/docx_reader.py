import docx
import os

def docx_to_text(docx_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        # Check if file exists
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
        
        # Open and process document
        doc = docx.Document(docx_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                paragraphs.append(para.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        text = "\n".join(paragraphs)
        return text
        
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {str(e)}")
